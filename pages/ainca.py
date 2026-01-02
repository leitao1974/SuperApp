import streamlit as st
import google.generativeai as genai
import pypdf
from docx import Document
from io import BytesIO
import time

import sys, os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import utils

# ... (depois do st.set_page_config) ...

utils.sidebar_comum()
# --- CONFIGURAÇÃO ---
st.set_page_config(page_title="Simplex AIncA AI Pro", page_icon="⚖️", layout="wide")
st.title("⚖️ Analista AIncA (Com Gestão de Quota)")

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("Configuração")
    api_key = st.text_input("Google API Key", type="password")
    
    # Forçar o modelo 1.5 Flash (é o mais estável para quotas)
    st.info("Modelo definido: gemini-1.5-flash (Melhor para documentos longos)")
    selected_model = "models/gemini-1.5-flash"

    if api_key:
        genai.configure(api_key=api_key)

    st.divider()
    st.markdown("""
    **Dicas para evitar erros de Quota:**
    1. Se o PDF do projeto for gigante (ex: 200 pág), tente carregar apenas a **Memória Descritiva**.
    2. Evite carregar peças desenhadas/plantas (a IA lê mal e ocupam muito espaço).
    """)

# --- FUNÇÕES ---
def get_text_from_multiple_files(file_list):
    combined_text = ""
    file_names = []
    if not file_list: return None, None

    for uploaded_file in file_list:
        try:
            reader = pypdf.PdfReader(uploaded_file)
            file_text = ""
            for page in reader.pages:
                file_text += page.extract_text() or "" 
            
            combined_text += f"\n--- FICHEIRO: {uploaded_file.name} ---\n{file_text}\n"
            file_names.append(uploaded_file.name)
        except Exception as e:
            st.error(f"Erro a ler {uploaded_file.name}: {e}")
            
    return combined_text, file_names

def create_word_docx(text, p_files, l_files):
    doc = Document()
    doc.add_heading('Relatório AIncA', 0)
    doc.add_paragraph(f"Projeto: {', '.join(p_files)}")
    doc.add_paragraph(f"Legislação: {', '.join(l_files)}")
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACE ---
col1, col2 = st.columns(2)
with col1:
    files_p = st.file_uploader("1. Projeto (Memória Descritiva)", type=["pdf"], accept_multiple_files=True)
with col2:
    files_l = st.file_uploader("2. Legislação (DL 11/2023, etc)", type=["pdf"], accept_multiple_files=True)

if 'analise' not in st.session_state: st.session_state['analise'] = None
if 'n_p' not in st.session_state: st.session_state['n_p'] = []
if 'n_l' not in st.session_state: st.session_state['n_l'] = []

# --- LÓGICA DE ANÁLISE ---
if st.button("🚀 Analisar", type="primary"):
    if not api_key or not files_p or not files_l:
        st.warning("Preencha a chave e carregue ficheiros.")
        st.stop()

    # 1. Leitura
    with st.status("A processar ficheiros...", expanded=True) as status:
        status.write("A extrair texto dos PDFs...")
        text_p, names_p = get_text_from_multiple_files(files_p)
        text_l, names_l = get_text_from_multiple_files(files_l)
        
        # 2. Verificação de Tamanho (PREVENÇÃO DE ERRO)
        total_chars = len(text_p) + len(text_l)
        estimated_tokens = total_chars / 4 # Estimativa grosseira
        status.write(f"Tamanho total detetado: ~{int(estimated_tokens)} tokens.")
        
        if estimated_tokens > 240000:
            status.update(label="❌ Ficheiros demasiado grandes!", state="error")
            st.error(f"Erro: O texto total tem ~{int(estimated_tokens)} tokens. O limite gratuito é 250.000. Por favor reduza o tamanho do PDF do Projeto (remova anexos ou plantas).")
            st.stop()

        # 3. Envio para IA com RETRY LOOP
        status.write("A contactar a Inteligência Artificial...")
        model = genai.GenerativeModel(selected_model)
        
        prompt = f"""Atua como Consultor Ambiental.
        LEGISLAÇÃO: {text_l}
        PROJETO: {text_p}
        TAREFA: Analisa a conformidade, dispensas Simplex e incumprimentos."""

        max_retries = 3
        wait_time = 40 # Segundos iniciais

        for attempt in range(max_retries):
            try:
                response = model.generate_content(prompt)
                st.session_state['analise'] = response.text
                st.session_state['n_p'] = names_p
                st.session_state['n_l'] = names_l
                status.update(label="✅ Análise Concluída!", state="complete")
                break # Sucesso, sai do loop

            except Exception as e:
                if "429" in str(e) or "Quota" in str(e):
                    if attempt < max_retries - 1:
                        status.write(f"⚠️ Limite de velocidade atingido. A aguardar {wait_time}s para tentar de novo (Tentativa {attempt+1}/{max_retries})...")
                        # Barra de progresso visual
                        progress_bar = st.progress(0)
                        for i in range(wait_time):
                            time.sleep(1)
                            progress_bar.progress((i + 1) / wait_time)
                        progress_bar.empty()
                        wait_time += 20 # Aumenta o tempo de espera na próxima
                    else:
                        status.update(label="❌ Falha Definitiva", state="error")
                        st.error("Não foi possível analisar mesmo após várias tentativas. Os ficheiros são demasiado pesados para o plano gratuito neste momento.")
                else:
                    status.update(label="❌ Erro Inesperado", state="error")
                    st.error(f"Erro: {e}")
                    break

# --- RESULTADOS ---
if st.session_state['analise']:
    st.divider()
    st.markdown("### 📋 Relatório")
    st.markdown(st.session_state['analise'])
    
    doc = create_word_docx(st.session_state['analise'], st.session_state['n_p'], st.session_state['n_l'])
    st.download_button("📥 Download Word", doc, "Relatorio_Simplex.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

