import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from io import BytesIO
from duckduckgo_search import DDGS
import requests
from bs4 import BeautifulSoup
import time

# Tenta importar a biblioteca legislativa local
try:
    import legislacao
except ImportError:
    st.error("⚠️ Ficheiro 'legislacao.py' não encontrado. Cria-o na mesma pasta do app.py.")
    st.stop()

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Análise Ambiental IA (Pro)",
    page_icon="🌿",
    layout="wide"
)

# --- GESTÃO DE ESTADO ---
if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0

def limpar_dados():
    st.session_state.uploader_key += 1
    st.rerun()

# --- ESTILO CSS ---
st.markdown("""
<style>
    .stButton>button { width: 100%; border-radius: 8px; height: 3em; font-weight: bold; }
    h1 { color: #155724; }
    .stExpander { border: 1px solid #c3e6cb; border-radius: 5px; background-color: #f8f9fa; }
    .stToast { background-color: #d4edda; color: #155724; }
</style>
""", unsafe_allow_html=True)

# --- CABEÇALHO ---
col1, col2 = st.columns([1, 6])
with col1: st.markdown("# 🌿")
with col2:
    st.title("Análise Ambiental & Compliance")
    st.caption("Protocolo PATE v5.0 | Full Context Window | Enterprise Ready")

# --- FUNÇÃO AUXILIAR: LISTAR MODELOS ---
def get_available_models(api_key):
    try:
        genai.configure(api_key=api_key)
        models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        return models
    except:
        return []

# --- SIDEBAR ---
with st.sidebar:
    st.header("⚙️ 1. Motor de IA")
    api_key = st.text_input("Google Gemini API Key", type="password")
    
    # Seletor de Modelo Inteligente
    selected_model = "models/gemini-1.5-flash"
    if api_key:
        avail = get_available_models(api_key)
        if avail:
            # Tenta selecionar o 1.5 Flash por defeito (melhor custo-benefício)
            idx = 0
            for i, m in enumerate(avail):
                if "1.5-flash" in m: 
                    idx = i
                    break
            selected_model = st.selectbox("Modelo:", avail, index=idx)
    
    st.divider()
    
    st.header("📚 2. Base Legislativa")
    library = legislacao.get_library()
    library_context = ""
    active_count = 0
    
    # Seletores de Legislação
    for category, laws in library.items():
        with st.expander(f"📂 {category}", expanded=False):
            for law_name, details in laws.items():
                if st.checkbox(law_name, value=False, key=law_name):
                    active_count += 1
                    library_context += f"- [ATIVA] {law_name} ({details['nivel']})\n  MANDATO: {details['mandato']}\n\n"
    
    if active_count > 0: st.success(f"✅ {active_count} regimes ativados.")

    st.divider()
    st.header("🌐 3. Fontes Extra")
    uploaded_legal_docs = st.file_uploader("PDFs Adicionais", type="pdf", accept_multiple_files=True, key=f"legal_{st.session_state.uploader_key}")
    search_query = st.text_input("Pesquisa Web")
    use_web_search = st.checkbox("Incluir Web", value=True)
    
    st.divider()
    if st.button("🗑️ Limpar Sessão"): limpar_dados()

# --- FUNÇÕES ---

def get_pdf_text(pdf_file):
    """Lê o PDF completo, sem limites de páginas."""
    text = ""
    try:
        reader = PdfReader(pdf_file)
        # REMOVIDO: Limite de páginas. Agora lê tudo.
        for page in reader.pages:
            text += page.extract_text() or ""
    except Exception as e:
        st.error(f"Erro ao ler {pdf_file.name}: {e}")
    return text

def search_online(query):
    if not query: return ""
    results_text = ""
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(f"{query} legislação oficial", max_results=3))
        for r in results:
            try:
                page = requests.get(r['href'], timeout=4)
                soup = BeautifulSoup(page.content, 'html.parser')
                # Apanha mais contexto da web
                text = "\n".join([p.text for p in soup.find_all('p')])[:4000]
                results_text += f"\n>>> WEB: {r['title']} <<<\n{text}\n"
            except: continue
        return results_text
    except: return ""

def create_docx(text):
    doc = Document()
    doc.add_heading('Relatório de Auditoria Ambiental', 0)
    for line in text.split('\n'):
        if line.startswith('# '): doc.add_heading(line[2:], 1)
        elif line.startswith('## '): doc.add_heading(line[3:], 2)
        elif line.startswith('### '): doc.add_heading(line[4:], 3)
        else: doc.add_paragraph(line.replace('*',''))
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b

def run_analysis(target_text, lib_ctx, manual_ctx, web_ctx, api_key, model_name):
    """Executa a análise com o modelo completo, sem cortes de texto."""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    # REMOVIDO: Limite de caracteres. 
    # O modelo pago aguenta >1M tokens, por isso enviamos tudo.
    
    full_context = ""
    if lib_ctx: full_context += f"\n=== BIBLIOTECA LEGISLATIVA ===\n{lib_ctx}"
    if manual_ctx: full_context += f"\n=== UPLOADS MANUAIS ===\n{manual_ctx}"
    if web_ctx: full_context += f"\n=== PESQUISA WEB ===\n{web_ctx}"

    prompt = f"""
    Atua como um **Consultor Especialista em Ambiente e Estratégia**.
    Realiza uma AUDITORIA TÉCNICA E LEGAL (Protocolo PATE) ao documento fornecido.

    --- BASE DE CONFORMIDADE (A TUA "VERDADE") ---
    {full_context}
    
    --- DOCUMENTO ALVO ---
    {target_text}
    
    ## INSTRUÇÕES DE ANÁLISE:
    
    1. **Resumo e Maturidade:** Identifica o objeto do plano/projeto e o seu estado de maturidade.
    
    2. **Check-up de Conformidade (Rigoroso):**
       - Cruza as medidas propostas com a Legislação fornecida.
       - Identifica conflitos com RAN, REN, Rede Natura 2000, ou metas climáticas (PNEC/Lei do Clima).
       - Se detetares omissões (ex: falta de referência a AIA), assinala como Risco.

    3. **Análise de Exequibilidade:**
       - Critica a qualidade dos dados de base (ex: proxies vs dados de campo).
       - Avalia a capacidade operacional e financeira proposta.

    4. **Matriz de Risco:**
       - Apresenta os principais riscos ambientais e legais por nível de gravidade.

    5. **Recomendações Práticas:**
       - 3 a 5 medidas corretivas imediatas ("Actionable Insights").
    
    Usa linguagem técnica, formal e cita as secções do documento analisado.
    """
    
    # Sistema de Retry (útil mesmo na versão paga para falhas de rede)
    for attempt in range(3):
        try:
            return model.generate_content(prompt).text
        except Exception as e:
            if "429" in str(e):
                time.sleep(5 * (attempt + 1)) # Espera curta
            else:
                return f"❌ Erro na API: {str(e)}"
    
    return "❌ Erro persistente de conexão."

# --- INTERFACE PRINCIPAL ---
st.subheader("📄 Documento Alvo")
uploaded_target = st.file_uploader("Carrega o Relatório/Plano", type="pdf", key=f"main_{st.session_state.uploader_key}")

if uploaded_target and api_key:
    if st.button("🚀 INICIAR ANÁLISE TOTAL", type="primary"):
        with st.spinner(f"A processar contexto completo com {selected_model}..."):
            
            # 1. Extração de Texto
            tgt_txt = get_pdf_text(uploaded_target)
            
            man_ctx = ""
            if uploaded_legal_docs:
                for f in uploaded_legal_docs: man_ctx += get_pdf_text(f)
            
            web_ctx = search_online(search_query) if use_web_search and search_query else ""
            
            # 2. Execução
            result = run_analysis(tgt_txt, library_context, man_ctx, web_ctx, api_key, selected_model)
            
            # 3. Output
            st.success("Análise concluída com sucesso!")
            
            col_res1, col_res2 = st.columns([1, 4])
            with col_res1:
                if st.button("🧹 Nova Análise"): limpar_dados()
            
            t1, t2 = st.tabs(["Relatório", "Exportar"])
            with t1: st.markdown(result)
            with t2:
                st.download_button("Descarregar Word (.docx)", create_docx(result), "Relatorio_Ambiental.docx")
                st.download_button("Descarregar Markdown (.md)", result, "Relatorio_Ambiental.md")

