import sys
import os
import re
import time
from io import BytesIO

# --- 1. CONFIGURAÇÃO DE CAMINHOS ---
current_dir = os.path.dirname(os.path.abspath(__file__))
root_dir = os.path.dirname(current_dir)
sys.path.insert(0, root_dir)

import utils
import streamlit as st
import google.generativeai as genai
import pypdf
from docx import Document
from docx.shared import Pt, RGBColor

# --- 2. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="AIncA (Rede Natura 2000)", 
    page_icon="🦅", 
    layout="wide"
)

# --- 3. BARRA LATERAL (Base) ---
try:
    utils.sidebar_comum()
except:
    pass

# --- 4. TÍTULO E ENQUADRAMENTO ---
st.title("🦅 Avaliação de Incidências Ambientais (AIncA)")
st.markdown("""
**Enquadramento Legal:** Decreto-Lei n.º 140/99, de 24 de abril (alterado pelos DL n.º 49/2005 e DL n.º 156-A/2013).

Este módulo gera um **Relatório Técnico Fundamentado**, cruzando evidências do projeto (com referência à página) com a legislação aplicável.
""")

# Recuperar API Key
api_key = st.session_state.get("api_key", "")
if not api_key:
    st.warning("⚠️ **Atenção:** API Key não detetada. Por favor insira-a no menu lateral esquerdo.")
    st.stop()

# ==========================================
# --- 5. BASE DE CONHECIMENTO SETORIAL ---
# ==========================================
SECTOR_GUIDES = {
    "Geral / Outros": "Guia da Comissão Europeia (2011) - Avaliação de planos e projetos.",
    "Infraestruturas Lineares (Estradas)": "Manual de apoio ICNB (2008) e Guia APA (2009) para Infraestruturas Rodoviárias.",
    "Linhas Elétricas (Transporte >110kV)": "Manual CIBIO/ICNF/REN (2020) - Muito Alta Tensão e Avifauna. Atenção a Áreas Críticas.",
    "Linhas Elétricas (Distribuição <110kV)": "Manual ICNB (2008) - Linhas de Distribuição e Avifauna.",
    "Parques Eólicos": "Guias ICNB (2008) para Morcegos e APA (2009) para Parques Eólicos.",
    "ETAR / Hidráulica": "Guia APA (2008) para ETARs.",
    "Indústria Extrativa": "Guia CCDR-LVT (2008) para Minas e Pedreiras."
}

# ==========================================
# --- FUNÇÕES ---
# ==========================================

def get_available_models(key):
    """Lista modelos disponíveis na API."""
    try:
        genai.configure(api_key=key)
        models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        return models
    except:
        # Fallback genérico se a API falhar a listagem
        return ["models/gemini-2.0-flash", "models/gemini-1.5-flash", "models/gemini-1.5-pro"]

def get_text_with_page_markers(file_list):
    """
    Extrai texto inserindo marcadores de página explícitos.
    Isso permite à IA citar: 'Conforme Pág. 12 do ficheiro X'.
    """
    combined_text = ""
    file_names = []
    if not file_list: return None, None

    for uploaded_file in file_list:
        try:
            reader = pypdf.PdfReader(uploaded_file)
            doc_name = uploaded_file.name
            
            combined_text += f"\n\n=== INÍCIO DO DOCUMENTO: {doc_name} ===\n"
            
            for i, page in enumerate(reader.pages):
                content = page.extract_text() or "[Página em branco ou imagem]"
                # INJEÇÃO DE METADADOS PARA A IA LER
                combined_text += f"\n[DOC: {doc_name} | PÁG. {i+1}]\n{content}\n"
            
            combined_text += f"=== FIM DO DOCUMENTO: {doc_name} ===\n"
            file_names.append(doc_name)
            
        except Exception as e:
            st.error(f"Erro a ler {uploaded_file.name}: {e}")
            
    return combined_text, file_names

def create_word_docx(text, p_files, l_files, tipologia):
    """Gera Word com formatação profissional."""
    doc = Document()
    
    # Estilo do Título
    title = doc.add_heading('Parecer Técnico AIncA Fundamentado', 0)
    title.alignment = 1 # Center
    
    # Metadados
    p = doc.add_paragraph()
    runner = p.add_run(f"Tipologia: {tipologia}\n")
    runner.bold = True
    p.add_run(f"Data da Análise: {time.strftime('%d/%m/%Y')}\n")
    p.add_run(f"Documentos Analisados: {', '.join(p_files) if p_files else 'N/A'}")
    
    doc.add_paragraph("---")
    
    # Processamento do Markdown para Word
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        
        if line.startswith('## '): 
            h = doc.add_heading(line.replace('##', '').strip(), 1)
            h.style.font.color.rgb = RGBColor(0, 51, 102) # Azul escuro
            
        elif line.startswith('### '): 
            h = doc.add_heading(line.replace('###', '').strip(), 2)
            
        elif line.startswith('- ') or line.startswith('* '): 
            p = doc.add_paragraph(style='List Bullet')
            # Tenta detetar citações [Doc X, Pag Y] e pôr a negrito
            parts = re.split(r'(\[.*?Pág.*?\])', line[2:], flags=re.IGNORECASE)
            for part in parts:
                run = p.add_run(part)
                if "[" in part and "Pág" in part:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(80, 80, 80) # Cinza escuro
                    
        elif line.startswith('>'): # Citações transcritas
            p = doc.add_paragraph(style='Intense Quote')
            p.add_run(line.replace('>', '').strip()).italic = True
            
        else: 
            doc.add_paragraph(line)
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_with_retry(model, prompt, max_retries=3):
    """Tenta gerar com gestão automática de erros de cota (429)."""
    wait_time = 15 
    for attempt in range(max_retries):
        try:
            return model.generate_content(prompt, request_options={"timeout": 600})
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "quota" in error_msg.lower():
                if attempt < max_retries - 1:
                    st.warning(f"⚠️ Cota momentânea atingida ({model.model_name}). Aguarde {wait_time}s para nova tentativa automática...")
                    time.sleep(wait_time)
                    wait_time += 15
                else:
                    raise e
            else:
                raise e

# ==========================================
# --- INTERFACE ---
# ==========================================

# --- A. BARRA LATERAL ---
with st.sidebar:
    st.divider()
    st.markdown("### 🧠 Motor de IA")
    
    opcoes_modelos = get_available_models(api_key)
    
    # --- LÓGICA DE PRIORIDADE REFORÇADA ---
    # Ordem de preferência: 2.5 Flash -> 2.0 Flash -> 1.5 Flash -> Qualquer Flash -> Outros
    priority_targets = ["2.5-flash", "2.0-flash", "1.5-flash", "flash"]
    idx_padrao = 0
    found = False
    
    for target in priority_targets:
        for i, m in enumerate(opcoes_modelos):
            if target in m.lower():
                idx_padrao = i
                found = True
                break
        if found: break
            
    selected_model = st.selectbox(
        "Modelo:", 
        opcoes_modelos, 
        index=idx_padrao,
        help="O sistema dá prioridade aos modelos Flash mais recentes para maior rapidez e eficiência."
    )
    
    st.divider()
    st.header("Contexto Setorial")
    selected_sector = st.selectbox("Setor:", list(SECTOR_GUIDES.keys()))
    st.info(f"📚 {SECTOR_GUIDES[selected_sector]}")

# --- B. UPLOADS ---
col1, col2 = st.columns(2)
with col1:
    files_p = st.file_uploader("1. Projeto (Obrigatório)", type=["pdf"], accept_multiple_files=True)
with col2:
    files_l = st.file_uploader("2. Cartografia/Anexos (Opcional)", type=["pdf"], accept_multiple_files=True)

# --- C. AÇÃO ---
if st.button("🚀 Gerar Relatório Fundamentado", type="primary", use_container_width=True):
    if not files_p:
        st.error("⚠️ Carregue os ficheiros do projeto.")
    else:
        with st.status("A realizar Auditoria AIncA...", expanded=True) as status:
            
            # 1. Leitura com Mapeamento de Páginas
            status.write("📖 A indexar páginas e documentos...")
            text_p, names_p = get_text_with_page_markers(files_p)
            text_l, names_l = get_text_with_page_markers(files_l)
            
            # 2. Configuração
            genai.configure(api_key=api_key)
            status.write(f"🤖 A analisar com **{selected_model}**...")
            model = genai.GenerativeModel(selected_model)
            
            guia_especifico = SECTOR_GUIDES[selected_sector]
            
            # 3. Prompt de Auditoria Rigorosa
            prompt = f"""
            Atua como Perito Sénior em Avaliação Ambiental (Especialista AIncA e Rede Natura 2000).
            A tua tarefa é produzir um RELATÓRIO TÉCNICO DE FUNDAMENTAÇÃO.

            === REGRAS DE OURO (OBRIGATÓRIAS) ===
            1. **CITAÇÃO DE FACTOS:** Qualquer afirmação sobre o projeto (distâncias, áreas, características) DEVE ter a fonte exata.
               Formato obrigatório: "O projeto ocupa 2ha..." [DOC: NomeDoFicheiro | PÁG. X].
            2. **TRANSCRIÇÃO:** Sempre que possível, transcreve pequenas frases do documento original entre aspas para provar o ponto.
               Ex: Como refere o promotor: "...não se preveem afetações..." [DOC: X | PÁG. Y].
            3. **FUNDAMENTAÇÃO LEGAL:** Cita sempre o artigo da lei aplicável (DL 140/99).

            === CONTEXTO TÉCNICO ===
            Setor: {selected_sector}
            Guia de Referência: {guia_especifico}
            
            === DADOS DO PROJETO (COM MARCADORES DE PÁGINA) ===
            {text_p}
            {text_l}
            
            === ESTRUTURA DO RELATÓRIO ===
            
            ## 1. DADOS DE IDENTIFICAÇÃO E ENQUADRAMENTO
            (Identifica o Promotor, Localização e Resumo do Projeto com base nos documentos. Cita a página da Memória Descritiva).
            
            ## 2. TRIAGEM JURÍDICA (SCREENING)
            - **Gestão do Sítio:** O projeto é para gestão da ZEC/ZPE? (Cita onde leste isto).
            - **Concorrência com AIA:** Verifica se o projeto cai nos Anexos do DL 151-B/2013. Se sim, conclui que a AIncA é integrada na AIA.
            - **Afetação Significativa:** Distância à Rede Natura 2000 mais próxima. Há sobreposição? [Cita Pág.]
            
            ## 3. ANÁLISE DE INCIDÊNCIAS (FACTOS E EVIDÊNCIAS)
            (Aqui deves usar as citações de página intensivamente).
            - Descritor Fauna/Flora: O que diz o projeto? [Cita Pág.]
            - Impactos na Integridade: O que diz o estudo de incidências? [Cita Pág.]
            - Cumprimento do Guia Setorial ({selected_sector}).
            
            ## 4. EVIDÊNCIAS TRANSCRITAS
            (Lista 3 a 5 frases chave copiadas ipsis verbis dos documentos que suportam a tua decisão).
            
            ## 5. CONCLUSÃO E PARECER TÉCNICO
            - O projeto carece de AIncA aprofundada?
            - Está dispensado?
            - Que medidas de mitigação são essenciais?
            """

            try:
                # 4. Geração
                response = generate_with_retry(model, prompt)
                
                status.update(label="✅ Relatório Gerado", state="complete")
                
                # Visualização
                st.markdown("### 🦅 Relatório Técnico Fundamentado")
                st.markdown(response.text)
                
                # Download
                doc = create_word_docx(response.text, names_p, names_l, selected_sector)
                st.download_button(
                    "📥 Descarregar Relatório (Word)", 
                    doc, 
                    "Relatorio_AIncA_Fundamentado.docx", 
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                status.update(label="❌ Erro", state="error")
                if "429" in str(e):
                    st.error(f"Cota excedida no modelo {selected_model}. Tente novamente em 1 minuto.")
                else:
                    st.error(f"Erro: {e}")
