import streamlit as st
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted
import io
import time
from datetime import datetime
import re

# ==========================================
# --- 1. BASE DE DADOS JURÍDICA ---
# ==========================================
LEGISLATION_DB = {
    "RJAIA (DL 151-B/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-116043164",
    "Alteração RJAIA (DL 152-B/2017)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/152-b-2017-114337069",
    "RGGR (DL 102-D/2020)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2020-150917243",
    "LUA (DL 75/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2015-106562356",
    "Rede Natura 2000 (DL 140/99)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/1999-34460975",
    "Regulamento Geral do Ruído (DL 9/2007)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2007-34526556",
    "Lei da Água (Lei 58/2005)": "https://diariodarepublica.pt/dr/legislacao-consolidada/lei/2005-34563267",
    "Emissões Industriais (DL 127/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-34789569"
}

# ==========================================
# --- CONFIGURAÇÃO INICIAL E ESTADO ---
# ==========================================
st.set_page_config(page_title="Análise Caso a Caso RJAIA", page_icon="⚖️", layout="wide")

if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0
if 'validation_result' not in st.session_state:
    st.session_state.validation_result = None
if 'decision_result' not in st.session_state:
    st.session_state.decision_result = None

def reset_app():
    st.session_state.uploader_key += 1
    st.session_state.validation_result = None
    st.session_state.decision_result = None

# ==========================================
# --- SIDEBAR & SETUP ---
# ==========================================
with st.sidebar:
    st.header("🔐 Configuração")
    
    api_key = None
    if "GOOGLE_API_KEY" in st.secrets:
        api_key = st.secrets["GOOGLE_API_KEY"]
        st.success("Chave API detetada (Secrets)!")
    else:
        api_key = st.text_input("Google API Key", type="password")
    
    selected_model = "gemini-1.5-flash"
    
    if api_key:
        try:
            genai.configure(api_key=api_key)
            models = genai.list_models()
            valid_models = [m.name for m in models if 'generateContent' in m.supported_generation_methods]
            
            if valid_models:
                idx = next((i for i, m in enumerate(valid_models) if 'flash' in m), 0)
                selected_model = st.selectbox("Modelo IA:", valid_models, index=idx)
                st.info("✅ Sistema Pronto")
            else:
                st.warning("⚠️ Chave válida, mas sem modelos disponíveis.")
        except Exception as e:
            st.error(f"Erro na API: {e}")
    
    st.divider()
    if st.button("🔄 Nova Análise / Limpar Tudo", use_container_width=True):
        reset_app()
        st.rerun()

# ==========================================
# --- FUNÇÕES AUXILIARES (WORD) ---
# ==========================================

def add_hyperlink(paragraph, text, url):
    """Adiciona um hiperlink clicável num parágrafo do Word."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def markdown_to_word(doc, text):
    """
    Converte Markdown para Word com formatação melhorada:
    - Espaçamento entre parágrafos
    - Listas com bullets reais
    - Negrito processado corretamente
    """
    if not text: return
    
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line: continue
        
        p = None
        
        # 1. Cabeçalhos
        if line.startswith('## '):
            p = doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line.replace('###', '').strip(), level=3)
        
        # 2. Listas (Bullets)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            clean_line = line[2:]
            process_bold(p, clean_line)
            # Adiciona um pequeno espaçamento após cada item de lista para não ficar tudo colado
            p.paragraph_format.space_after = Pt(6) 
            
        # 3. Parágrafos Normais
        else:
            p = doc.add_paragraph()
            process_bold(p, line)
            # Adiciona espaçamento padrão após parágrafos
            p.paragraph_format.space_after = Pt(10)
        
        # Justificação para texto normal (não cabeçalhos)
        if p and not line.startswith('#'):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def process_bold(paragraph, text):
    """Processa negrito (**texto**) dentro de parágrafos."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def append_legislation_section(doc):
    doc.add_page_break()
    doc.add_heading("Legislação Consultada e Referências", level=1)
    
    p_intro = doc.add_paragraph("A presente análise teve por base os seguintes diplomas legais:")
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for name, url in LEGISLATION_DB.items():
        p = doc.add_paragraph(style='List Bullet')
        add_hyperlink(p, name, url)

# ==========================================
# --- EXTRAÇÃO E IA ---
# ==========================================

def extract_text(files, label):
    text = ""
    if not files: return "" 
    
    for f in files:
        try:
            f.seek(0)
            bytes_data = f.read()
            f_stream = io.BytesIO(bytes_data)
            r = PdfReader(f_stream)
            
            if r.is_encrypted:
                try: r.decrypt("") 
                except: pass

            file_text = ""
            for i, p in enumerate(r.pages):
                extracted = p.extract_text()
                if extracted:
                    file_text += f"[Pág. {i+1}] {extracted}\n"
            
            if len(file_text.strip()) < 50:
                st.warning(f"⚠️ O ficheiro '{f.name}' parece ser uma imagem ou está vazio.")
                text += f"\n[AVISO: Ficheiro {f.name} ilegível/imagem]\n"
            else:
                text += f"\n\n>>> FONTE: {label} ({f.name}) <<<\n{file_text}"
                
        except Exception as e:
            st.error(f"❌ Erro ao ler '{f.name}': {str(e)}")
            text += f"\n[ERRO LEITURA: {f.name}]\n"
            
    return text

def get_ai(prompt):
    if not api_key: return "Erro: Falta API Key."
    model = genai.GenerativeModel(selected_model)
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Erro IA: {str(e)}"

# --- PROMPTS MELHORADOS (ESTRUTURA FORÇADA) ---

def analyze_validation(t_sim, t_form, t_proj, t_leg):
    legislacao_str = ", ".join(LEGISLATION_DB.keys())
    return get_ai(f"""
    Atua como PERITO AUDITOR AMBIENTAL rigoroso.
    
    CONTEXTO LEGAL:
    Utiliza: {legislacao_str}.
    Contexto Local Prioritário: {t_leg[:30000]}

    DADOS:
    {t_sim[:25000]}
    {t_form[:25000]}
    {t_proj[:80000]}

    TAREFA:
    Audita a consistência, valida limiares RJAIA e cruza com PDM.

    REGRAS DE OUTPUT (MARKDOWN OBRIGATÓRIO):
    1. Usa cabeçalhos "##" para secções principais.
    2. Usa cabeçalhos "###" para sub-secções.
    3. Usa LISTAS COM MARCADORES (-) para apresentar dados. Não uses parágrafos gigantes.
    4. Usa **Negrito** para destacar valores numéricos e conclusões (Ex: **VALIDADO**, **INCONSISTENTE**).

    ESTRUTURA DO RELATÓRIO:
    Linha 1: "STATUS: [VALIDADO ou INCONSISTENTE]"
    
    ## 1. Resumo Executivo
    (Breve síntese do projeto e conclusão principal).

    ## 2. Auditoria de Consistência
    ### 2.1. Áreas
    - **Simulação:** [Valor]
    - **Projeto:** [Valor]
    - **Conclusão:** [Análise]

    ### 2.2. Gestão de Resíduos e Capacidades
    - **LER Previstos:** [Lista]
    - **Capacidade Tratamento:** [Valor] vs [Valor]
    - **Capacidade Armazenamento:** [Valor] vs [Valor]

    ## 3. Enquadramento Legal e Localização
    ### 3.1. Análise RJAIA
    - **Limiar Aplicável:** [Citar Anexo/Ponto]
    - **Valor do Projeto:** [Valor]
    - **Parecer:** [Sujeito / Não Sujeito]

    ### 3.2. Condicionantes Locais (PDM/REN/RAN)
    - **Localização:** [Freguesia/Local]
    - **Análise PDM:** (Cruzar com o documento 'Contexto Local' se existir).
    - **Outras Condicionantes:** (REN, RAN, Recursos Hídricos).
    """)

def generate_decision_text(t_sim, t_form, t_proj, t_leg):
    return get_ai(f"""
    Atua como Técnico Superior da CCDR. Redige a MINUTA DE DECISÃO.
    
    DADOS:
    {t_proj[:120000]}
    {t_form[:25000]}
    Legislação Local: {t_leg[:40000]}

    OUTPUT - APENAS PREENCHER AS TAGS ABAIXO (Texto formal e justificado):
    
    ### CAMPO_DESIGNACAO
    (Nome exato do projeto)
    
    ### CAMPO_TIPOLOGIA
    (Ex: Ponto 11.b do Anexo II do DL 151-B/2013...)
    
    ### CAMPO_ENQUADRAMENTO
    (Artigo jurídico que fundamenta a decisão 'Caso a Caso')
    
    ### CAMPO_LOCALIZACAO
    (Localização administrativa completa)
    
    ### CAMPO_AREAS_SENSIVEIS
    (Análise art. 2.º RJAIA. Se não houver, escrever "Não abrange áreas sensíveis.")
    
    ### CAMPO_PROPONENTE
    (Nome da empresa/promotor)
    
    ### CAMPO_ENTIDADE_LICENCIADORA
    (Câmara Municipal ou CCDR, conforme aplicável)
    
    ### CAMPO_AUTORIDADE_AIA
    (CCDR territorialmente competente)
    
    ### CAMPO_DESCRICAO
    (Descrição técnica. Usa parágrafos curtos.)
    
    ### CAMPO_CARATERISTICAS
    (Detalhar capacidades instaladas, LER, gestão de efluentes. Usa linguagem técnica.)
    
    ### CAMPO_LOCALIZACAO_PROJETO
    (Análise da compatibilidade com o PDM e servidões. Cita a 'Legislação Local' se fornecida.)
    
    ### CAMPO_IMPACTES
    (Análise por descritor: Ruído, Ar, Água, Resíduos. Separa por parágrafos.)
    
    ### CAMPO_DECISAO
    (Decisão final explícita: SUJEITO / NÃO SUJEITO a AIA)
    
    ### CAMPO_CONDICIONANTES
    (Lista de obrigações a cumprir pelo promotor.)
    """)

# ==========================================
# --- GERADORES DE DOCS ---
# ==========================================

def create_validation_doc(text):
    doc = Document()
    
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "Relatório de Auditoria Técnica"
    sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    h = doc.add_heading("Auditoria de Conformidade Legal e Técnica", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    if text:
        p_status = doc.add_paragraph()
        p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_text = "PARECER: ANÁLISE CONCLUÍDA"
        color = RGBColor(0, 0, 0)
        
        if "INCONSISTENTE" in text.upper():
            status_text = "⚠️ PARECER: INCONGRUÊNCIAS DETETADAS"
            color = RGBColor(255, 0, 0)
        elif "VALIDADO" in text.upper():
            status_text = "✅ PARECER: DADOS CONSISTENTES"
            color = RGBColor(0, 128, 0)
            
        r = p_status.add_run(status_text)
        r.font.color.rgb = color
        r.bold = True
        r.font.size = Pt(14)
        
        doc.add_paragraph("---")
        clean_text = re.sub(r'STATUS:.*', '', text, count=1).strip()
        markdown_to_word(doc, clean_text)
    else:
        doc.add_paragraph("Erro: Sem conteúdo gerado.")
    
    append_legislation_section(doc)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def create_decision_doc(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6) 
    
    if not text: return io.BytesIO()

    def get_tag(tag):
        m = re.search(f"### {tag}(.*?)###", text, re.DOTALL)
        if not m: m = re.search(f"### {tag}(.*)", text, re.DOTALL)
        return m.group(1).strip() if m else "..."

    h = doc.add_heading("Análise prévia e decisão de sujeição a AIA", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_merged_header(txt):
        r = table.add_row()
        c = r.cells[0]
        c.merge(r.cells[1])
        p = c.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        c.paragraphs[0].paragraph_format.space_after = Pt(6)
        return r

    def add_row(label, val):
        r = table.add_row()
        r.cells[0].paragraphs[0].add_run(label).bold = True
        r.cells[1].paragraphs[0].text = val
        return r

    def add_full_text(header, content):
        add_merged_header(header)
        r = table.add_row()
        c = r.cells[0]
        c.merge(r.cells[1])
        c.paragraphs[0].text = "" # Limpar
        
        # Processar parágrafos
        paras = content.split('\n')
        for para in paras:
            if para.strip():
                p = c.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)

    add_merged_header("Identificação")
    add_row("Designação do projeto", get_tag("CAMPO_DESIGNACAO"))
    add_row("Tipologia", get_tag("CAMPO_TIPOLOGIA"))
    add_row("Enquadramento", get_tag("CAMPO_ENQUADRAMENTO"))
    add_row("Localização", get_tag("CAMPO_LOCALIZACAO"))
    add_row("Áreas Sensíveis", get_tag("CAMPO_AREAS_SENSIVEIS"))
    add_row("Proponente", get_tag("CAMPO_PROPONENTE"))
    
    add_full_text("Descrição", get_tag("CAMPO_DESCRICAO"))
    add_merged_header("Fundamentação")
    add_full_text("Caraterísticas", get_tag("CAMPO_CARATERISTICAS"))
    add_full_text("Localização", get_tag("CAMPO_LOCALIZACAO_PROJETO"))
    add_full_text("Impactes", get_tag("CAMPO_IMPACTES"))

    add_merged_header("Decisão")
    r = table.add_row()
    c = r.cells[0]
    c.merge(r.cells[1])
    p = c.paragraphs[0]
    run = p.add_run(get_tag("CAMPO_DECISAO"))
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_full_text("Condicionantes", get_tag("CAMPO_CONDICIONANTES"))

    doc.add_paragraph("\n")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("O Técnico,\n\n_______________________").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# ==========================================
# --- UI PRINCIPAL ---
# ==========================================
st.title("⚖️ Análise Caso a Caso (RJAIA)")
st.markdown("### Auditoria Técnica e Decisão Fundamentada")

col1, col2, col3, col4 = st.columns(4)
with col1: files_sim = st.file_uploader("📂 Simulação SILiAmb", type=['pdf'], accept_multiple_files=True, key=f"s_{st.session_state.uploader_key}")
with col2: files_form = st.file_uploader("📂 Formulário", type=['pdf'], accept_multiple_files=True, key=f"f_{st.session_state.uploader_key}")
with col3: files_doc = st.file_uploader("📂 Projeto/Memória", type=['pdf'], accept_multiple_files=True, key=f"p_{st.session_state.uploader_key}")
with col4: files_leg = st.file_uploader("📜 Legislação Local", type=['pdf'], accept_multiple_files=True, key=f"l_{st.session_state.uploader_key}")

st.markdown("---")

if st.button("🚀 Processar", type="primary", use_container_width=True):
    if not (files_sim and files_form and files_doc):
        st.error("Carregue Simulação, Formulário e Projeto.")
    elif not api_key:
        st.error("API Key em falta.")
    else:
        try:
            with st.status("⚙️ A processar...", expanded=True) as status:
                st.write("📄 Lendo documentos...")
                ts = extract_text(files_sim, "SIM")
                tf = extract_text(files_form, "FORM")
                tp = extract_text(files_doc, "PROJ")
                tl = extract_text(files_leg, "LOCAL") if files_leg else "N/A"
                
                st.write("🕵️ Auditoria (Modo Estruturado)...")
                val = analyze_validation(ts, tf, tp, tl)
                st.session_state.validation_result = val
                
                st.write("⚖️ Decisão...")
                dec = generate_decision_text(ts, tf, tp, tl)
                st.session_state.decision_result = dec
                
                status.update(label="Concluído!", state="complete")
        except Exception as e:
            st.error(f"Erro: {e}")

if st.session_state.validation_result:
    c1, c2 = st.columns(2)
    f_val = create_validation_doc(st.session_state.validation_result)
    c1.download_button("📄 Relatório Auditoria", f_val.getvalue(), "Auditoria.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    f_dec = create_decision_doc(st.session_state.decision_result)
    c2.download_button("📝 Minuta Decisão", f_dec.getvalue(), "Decisao.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
