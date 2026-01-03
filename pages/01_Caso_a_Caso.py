import sys
import os

# --- 1. CONFIGURAÇÃO DE CAMINHOS ---
# Garante que o Python encontra o utils.py na pasta raiz
current_dir = os.path.dirname(os.path.abspath(__file__))
root_dir = os.path.dirname(current_dir)
sys.path.insert(0, root_dir)

import utils
import streamlit as st
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai
import io
import re
from datetime import datetime

# --- 2. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Análise Caso a Caso",
    page_icon="⚖️",
    layout="wide"
)

# --- 3. BARRA LATERAL (Menu e Chave) ---
try:
    utils.sidebar_comum()
except:
    pass

# --- 4. TÍTULO E VERIFICAÇÃO DE CHAVE ---
st.title("⚖️ Análise Caso a Caso (RJAIA)")
st.markdown("### Auditoria Técnica e Decisão Fundamentada")
st.caption("Verificação de critérios de sujeição a AIA (Anexo II do DL 151-B/2013).")

# Recupera chave da memória
api_key = st.session_state.get("api_key", "")
if not api_key:
    st.warning("⚠️ **Atenção:** API Key não detetada. Por favor insira-a no menu lateral esquerdo.")
    st.stop()

# ==========================================
# --- 5. SELETOR DE MODELO (DINÂMICO) ---
# ==========================================
def get_available_models(key):
    """Lista modelos disponíveis na API."""
    try:
        genai.configure(api_key=key)
        return [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    except:
        return ["models/gemini-2.0-flash", "models/gemini-1.5-flash"]

with st.sidebar:
    st.divider()
    st.markdown("### 🧠 Motor de IA")
    
    opcoes_modelos = get_available_models(api_key)
    
    # Lógica de Prioridade: 2.5 Flash > 2.0 Flash > 1.5 Flash > Qualquer Flash
    targets = ["2.5-flash", "2.0-flash", "1.5-flash", "flash"]
    idx_padrao = 0
    found = False
    
    for t in targets:
        for i, m in enumerate(opcoes_modelos):
            if t in m.lower():
                idx_padrao = i
                found = True
                break
        if found: break
            
    selected_model = st.selectbox(
        "Modelo:", 
        opcoes_modelos, 
        index=idx_padrao,
        help="O sistema seleciona automaticamente o modelo Flash mais recente para maior rapidez."
    )

# ==========================================
# --- 6. FUNÇÕES AUXILIARES ---
# ==========================================

LEGISLATION_DB = {
    "RJAIA (DL 151-B/2013)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2013-116043164",
    "Alteração RJAIA (DL 152-B/2017)": "https://diariodarepublica.pt/dr/detalhe/decreto-lei/152-b-2017-114337069",
    "LUA (DL 75/2015)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/2015-106562356",
    "Rede Natura 2000 (DL 140/99)": "https://diariodarepublica.pt/dr/legislacao-consolidada/decreto-lei/1999-34460975"
}

def extract_text(files, label):
    """Extrai texto de PDFs carregados."""
    text = ""
    if not files: return "" 
    for f in files:
        try:
            reader = PdfReader(f)
            file_text = ""
            for page in reader.pages:
                file_text += page.extract_text() or ""
            text += f"\n\n>>> FONTE: {label} ({f.name}) <<<\n{file_text}"
        except Exception as e:
            st.error(f"Erro ao ler '{f.name}': {e}")
    return text

def analyze_validation(t_sim, t_form, t_proj, t_leg, key, model_name):
    """Executa a auditoria de validação (Análise Técnica)."""
    genai.configure(api_key=key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Atua como Auditor Ambiental da Autoridade de AIA.
    
    CONTEXTO LEGAL:
    Utiliza: RJAIA (DL 151-B/2013) e seus Anexos.
    Contexto Local (PDM/Condicionantes): {t_leg[:20000]}

    DADOS DO PROJETO:
    SIMULAÇÃO: {t_sim[:20000]}
    FORMULÁRIO: {t_form[:20000]}
    MEMÓRIA DESCRITIVA: {t_proj[:60000]}

    TAREFA:
    Realiza uma auditoria de conformidade para verificar se o projeto está bem enquadrado como "Caso a Caso" ou se devia ser AIA direta.
    Verifica limiares do Anexo II.

    OUTPUT (MARKDOWN):
    ## 1. Resumo do Projeto
    ## 2. Verificação de Limiares (Anexo II)
    - Ponto do Anexo: [Identificar]
    - Limiar Legal: [Valor]
    - Valor do Projeto: [Valor]
    - Parecer: [Cumpre/Não Cumpre]
    ## 3. Análise de Sensibilidade (Localização)
    ## 4. Conclusão da Validação
    """
    return model.generate_content(prompt).text

def generate_decision_text(t_sim, t_form, t_proj, t_leg, key, model_name):
    """Gera a Minuta de Decisão Final."""
    genai.configure(api_key=key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Redige a MINUTA DE DECISÃO FINAL (Técnico Superior).
    
    DADOS: {t_proj[:80000]} {t_form[:20000]}
    
    OUTPUT - APENAS O TEXTO PARA PREENCHER O WORD (Não uses Markdown aqui, usa texto corrido estruturado):
    
    Identificação do Projeto: [Nome]
    Promotor: [Nome]
    Localização: [Local]
    
    CONSIDERANDO QUE:
    1. O projeto se enquadra na alínea [X] do ponto [Y] do Anexo II...
    2. Da análise efetuada, verifica-se que [Resumo dos impactes]...
    3. Foram consultadas as entidades [Entidades]...
    
    DECISÃO:
    Face ao exposto, propõe-se a [SUJEIÇÃO / NÃO SUJEIÇÃO] a Avaliação de Impacte Ambiental.
    
    CONDICIONANTES:
    (Lista de condicionantes a cumprir caso não seja sujeito a AIA).
    """
    return model.generate_content(prompt).text

def create_doc_from_text(text, title):
    """Gera um ficheiro Word simples."""
    doc = Document()
    doc.add_heading(title, 0)
    
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('## '): doc.add_heading(line.replace('##', ''), 1)
        elif line.startswith('- '): doc.add_paragraph(line[2:], style='List Bullet')
        else: doc.add_paragraph(line)
        
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==========================================
# --- 7. INTERFACE PRINCIPAL ---
# ==========================================

# Gestão de Estado da Sessão
if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0
if 'validation_result' not in st.session_state: st.session_state.validation_result = None
if 'decision_result' not in st.session_state: st.session_state.decision_result = None

def reset_app():
    st.session_state.uploader_key += 1
    st.session_state.validation_result = None
    st.session_state.decision_result = None
    st.rerun()

# Uploads
c1, c2, c3, c4 = st.columns(4)
with c1: 
    files_sim = st.file_uploader("📂 Simulação SILiAmb", type=['pdf'], accept_multiple_files=True, key=f"s_{st.session_state.uploader_key}")
with c2: 
    files_form = st.file_uploader("📂 Formulário", type=['pdf'], accept_multiple_files=True, key=f"f_{st.session_state.uploader_key}")
with c3: 
    files_doc = st.file_uploader("📂 Projeto/Memória", type=['pdf'], accept_multiple_files=True, key=f"p_{st.session_state.uploader_key}")
with c4: 
    files_leg = st.file_uploader("📜 Legislação/PDM", type=['pdf'], accept_multiple_files=True, key=f"l_{st.session_state.uploader_key}")

st.markdown("---")

# Botão de Ação
if st.button("🚀 Processar Análise", type="primary", use_container_width=True):
    if not (files_sim and files_form and files_doc):
        st.error("⚠️ Faltam ficheiros obrigatórios (Simulação, Formulário e Projeto).")
    elif not api_key:
        st.error("🛑 API Key em falta.")
    else:
        with st.status("⚙️ A processar...", expanded=True) as status:
            st.write("📖 A ler documentos...")
            ts = extract_text(files_sim, "SIM")
            tf = extract_text(files_form, "FORM")
            tp = extract_text(files_doc, "PROJ")
            tl = extract_text(files_leg, "LOCAL") if files_leg else "N/A"
            
            # Validação
            st.write(f"🕵️ A realizar Auditoria Técnica com **{selected_model}**...")
            val = analyze_validation(ts, tf, tp, tl, api_key, selected_model)
            st.session_state.validation_result = val
            
            # Decisão
            st.write("⚖️ A redigir Minuta de Decisão...")
            dec = generate_decision_text(ts, tf, tp, tl, api_key, selected_model)
            st.session_state.decision_result = dec
            
            status.update(label="✅ Concluído!", state="complete")

# Área de Resultados
if st.session_state.validation_result:
    st.divider()
    st.success("Processo concluído com sucesso!")
    
    col_res1, col_res2 = st.columns(2)
    
    with col_res1:
        st.subheader("📄 Relatório de Auditoria")
        with st.expander("Ver Pré-visualização", expanded=False):
            st.markdown(st.session_state.validation_result)
        
        f_val = create_doc_from_text(st.session_state.validation_result, "Relatório de Auditoria Técnica")
        st.download_button(
            "📥 Descarregar Auditoria (.docx)", 
            f_val, 
            "Auditoria_Caso_a_Caso.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    with col_res2:
        st.subheader("📝 Minuta de Decisão")
        with st.expander("Ver Pré-visualização", expanded=False):
            st.text(st.session_state.decision_result) # Usa text para monospaced
            
        f_dec = create_doc_from_text(st.session_state.decision_result, "Minuta de Decisão")
        st.download_button(
            "📥 Descarregar Decisão (.docx)", 
            f_dec, 
            "Decisao_Final.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
    
    st.divider()
    if st.button("🔄 Limpar e Começar de Novo"):
        reset_app()

