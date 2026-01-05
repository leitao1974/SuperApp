import sys
import os
import re
import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from duckduckgo_search import DDGS
import time

# --- 1. CONFIGURAÇÃO DE CAMINHOS ---
current_dir = os.path.dirname(os.path.abspath(__file__))
root_dir = os.path.dirname(current_dir)
sys.path.insert(0, root_dir)

try:
    import utils
    try:
        import legislacao
    except ImportError:
        legislacao = None
except ImportError:
    pass

# --- 2. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Auditoria Académica & Compliance",
    page_icon="🎓",
    layout="wide"
)

# --- 3. TÍTULO E CHAVE ---
st.title("🎓 Auditoria Ambiental: Análise Académica")
st.markdown("""
**Protocolo PATE (Fundamentação Técnica e Científica).**
Gera relatórios com rigor académico, citação de fontes e análise detalhada de indicadores.
""")

api_key = st.session_state.get("api_key", "")
if not api_key:
    st.warning("⚠️ **Atenção:** API Key não detetada. Por favor insira-a no menu lateral esquerdo.")
    st.stop()

# ==========================================
# --- 4. FUNÇÕES ---
# ==========================================

def get_available_models(key):
    """Lista os modelos disponíveis na API."""
    try:
        genai.configure(api_key=key)
        models = genai.list_models()
        return [m.name for m in models if 'generateContent' in m.supported_generation_methods]
    except:
        return ["models/gemini-2.0-flash", "models/gemini-1.5-flash"]

def get_pdf_text_with_pages(pdf_file, simple_citation=False):
    """Extrai texto com paginação para citação académica."""
    text = ""
    try:
        reader = PdfReader(pdf_file)
        doc_name = pdf_file.name
        
        text += f"\n\n=== DOCUMENTO FONTE: {doc_name} ===\n"
        for i, page in enumerate(reader.pages):
            content = page.extract_text() or "[Página em branco/imagem]"
            citation = f"[PÁG. {i+1}]" if simple_citation else f"[DOC: {doc_name} | PÁG. {i+1}]"
            text += f"\n{citation}\n{content}\n"
        text += f"=== FIM DE: {doc_name} ===\n"
    except Exception as e:
        st.error(f"Erro ao ler PDF {pdf_file.name}: {e}")
    return text

def search_online(query):
    """Pesquisa Web para contexto complementar."""
    if not query: return ""
    results_text = ""
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(f"{query} legislação portugal ecologia", max_results=3))
        for r in results:
            results_text += f"\n>>> FONTE EXTERNA (WEB): {r['title']} ({r['href']}) <<<\n{r['body']}\n"
        return results_text
    except Exception as e:
        return f"Erro na pesquisa web: {str(e)}"

def format_paragraph(paragraph, text):
    """
    Processa o texto para o Word:
    1. Interpreta Markdown Bold (**texto**) e converte para Word Bold.
    2. Formata citações [DOC... | PÁG...] em cinza/negrito.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_text = part.replace('**', '')
            run = paragraph.add_run(clean_text)
            run.bold = True
        else:
            citation_parts = re.split(r'(\[.*?PÁG.*?\])', part)
            for sub_part in citation_parts:
                run = paragraph.add_run(sub_part)
                if "[" in sub_part and "PÁG" in sub_part and "]" in sub_part:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(80, 80, 80) # Cinza escuro
                    run.bold = True

def create_docx(text):
    """Gera DOCX com formatação limpa e justificada."""
    doc = Document()
    
    title = doc.add_heading('Parecer Técnico de Auditoria Ambiental', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_date = doc.add_paragraph(f"Data da Emissão: {time.strftime('%d/%m/%Y')}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("---")
    
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        
        if line.startswith('## '): 
            clean_line = line.replace('## ', '').replace('**', '') 
            h = doc.add_heading(clean_line, level=1)
            h.style.font.color.rgb = RGBColor(0, 50, 100) 
            
        elif line.startswith('### '): 
            clean_line = line.replace('### ', '').replace('**', '')
            doc.add_heading(clean_line, level=2)
            
        elif line.startswith('- ') or line.startswith('* '): 
            clean_line = line[2:] 
            p = doc.add_paragraph(style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            format_paragraph(p, clean_line) 
            
        elif line.startswith('>'): 
            p = doc.add_paragraph(style='Intense Quote')
            clean_line = line.replace('>', '').strip()
            p.add_run(clean_line).italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        else: 
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            format_paragraph(p, line) 
            
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b

def run_analysis(target_text, lib_ctx, manual_ctx, web_ctx, key, model_name):
    genai.configure(api_key=key)
    model = genai.GenerativeModel(model_name)
    
    # ATUALIZAÇÃO DO PROMPT: Inclusão de regra para Indicadores
    prompt = f"""
    Atua como **Auditor Ambiental Sénior e Investigador Académico**.
    
    === CONTEXTO LEGAL (Legislatura) ===
    {lib_ctx}
    
    === ANEXOS TÉCNICOS ===
    {manual_ctx}
    
    === PESQUISA RECENTE ===
    {web_ctx}
    
    === DOCUMENTO EM ANÁLISE ===
    {target_text}
    
    TAREFA:
    Elaborar um **Parecer Técnico de Auditoria** com elevado rigor científico.
    
    REGRA DE OURO (INDICADORES):
    Se o documento apresentar indicadores de desempenho, monitorização ou qualidade (KPIs), estes DEVEM ser tratados num capítulo próprio. Deves extrair a designação, a meta/valor de referência e fazer uma análise crítica sobre a sua pertinência ou dados apresentados.
    
    DIRETRIZES DE ESTILO:
    1.  **Impessoalidade Académica:** Usa a 3.ª pessoa (ex: "Verifica-se", "Constata-se").
    2.  **Fundamentação:** Todas as afirmações devem ser sustentadas por evidências textuais [CITAR].
    3.  **Formatação:** Usa **negrito** para destacar conceitos chave.
    
    ESTRUTURA DO PARECER:
    
    ## 1. Enquadramento e Análise de Maturidade
    (Síntese técnica do objeto de estudo e estado da arte do projeto).
    
    ## 2. Verificação de Conformidade Legal e Normativa
    (Análise comparativa entre o projeto e a legislação).
    - **[Diploma/Norma]:** [Análise] -> Evidência: "..." [CITAR].
    
    ## 3. Análise de Indicadores e Monitorização (KPIs)
    (Obrigatório se existirem indicadores. Caso contrário, indicar que não foram identificados).
    - **[Nome do Indicador]:** [Descrição Resumida / Meta].
      Análise Crítica: [Avaliar robustez dos dados, baseline ou tendência] [CITAR].
    
    ## 4. Identificação de Riscos Críticos e Lacunas
    (Diagnóstico de omissões, falhas metodológicas ou ausência de dados).
    
    ## 5. Conclusões e Recomendações Técnicas
    (Síntese conclusiva e medidas corretivas numeradas).
    """
    
    try:
        return model.generate_content(prompt, request_options={"timeout": 600}).text
    except Exception as e:
        return f"Erro na análise: {e}"

# ==========================================
# --- 5. INTERFACE ---
# ==========================================

if 'uploader_key' not in st.session_state: st.session_state.uploader_key = 0

# --- A. BARRA LATERAL ---
with st.sidebar:
    try:
        utils.sidebar_comum()
    except:
        pass
    st.divider()
    st.markdown("### 🧠 Motor de Inferência")
    opcoes_modelos = get_available_models(api_key)
    
    idx_padrao = 0
    targets = ["2.5-flash", "2.0-flash", "1.5-flash", "pro"]
    for t in targets:
        for i, m in enumerate(opcoes_modelos):
            if t in m.lower():
                idx_padrao = i
                break
        else: continue
        break
            
    selected_model = st.selectbox("Modelo:", opcoes_modelos, index=idx_padrao)

# --- B. ÁREA PRINCIPAL ---
library = legislacao.get_library() if legislacao else {}
lib_context = ""

with st.expander("📚 Base Legislativa (Referências)", expanded=False):
    c1, c2 = st.columns(2)
    i = 0
    for cat, laws in library.items():
        with (c1 if i % 2 == 0 else c2):
            st.markdown(f"**{cat}**")
            for name, info in laws.items():
                if st.checkbox(name, key=f"leg_{name}"):
                    desc = info.get('mandato', info.get('descricao', 'N/A'))
                    lib_context += f"- {name}: {desc}\n"
        i += 1

col_main, col_extra = st.columns([3, 2])
with col_main:
    st.subheader("📄 Documento Objeto de Análise")
    f_main = st.file_uploader("Relatório/Projeto (PDF)", type="pdf", key=f"main_{st.session_state.uploader_key}")

with col_extra:
    st.subheader("🔗 Elementos Complementares")
    f_extra = st.file_uploader("Anexos/Legislação (PDF)", type="pdf", accept_multiple_files=True, key=f"extra_{st.session_state.uploader_key}")
    web_q = st.text_input("Pesquisa Bibliográfica/Web", help="Ex: 'Regulamento UE 2024/1991 artigo 12'")

if st.button("⚖️ EMITIR PARECER TÉCNICO", type="primary", use_container_width=True):
    if not f_main:
        st.warning("⚠️ É necessário submeter o documento principal.")
    else:
        with st.status("⚙️ A processar auditoria académica...", expanded=True):
            
            # Leitura Inteligente
            has_extras = True if f_extra else False
            st.write("📖 A analisar corpus documental...")
            
            txt_main = get_pdf_text_with_pages(f_main, simple_citation=not has_extras)
            
            txt_extra = ""
            if f_extra:
                for f in f_extra: 
                    txt_extra += get_pdf_text_with_pages(f, simple_citation=False) + "\n"
            
            txt_web = ""
            if web_q:
                st.write("🌍 A consultar fontes externas...")
                txt_web = search_online(web_q)
            
            # Análise
            st.write(f"🧠 A elaborar parecer com **{selected_model}**...")
            res = run_analysis(txt_main, lib_context, txt_extra, txt_web, api_key, selected_model)
            
            # Output
            st.success("Parecer emitido com sucesso.")
            st.markdown("### 📝 Parecer Técnico")
            st.markdown(res)
            
            st.download_button(
                "📥 Descarregar Parecer (DOCX)", 
                create_docx(res), 
                "Parecer_Tecnico_Ambiental.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
