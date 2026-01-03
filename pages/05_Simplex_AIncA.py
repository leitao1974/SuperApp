import sys
import os

# --- 1. LIGAÇÃO AO UTILS (CRÍTICO) ---
# Isto garante que encontramos o ficheiro 'utils.py' na pasta de trás
current_dir = os.path.dirname(os.path.abspath(__file__))
root_dir = os.path.dirname(current_dir)
sys.path.insert(0, root_dir)

import streamlit as st
import utils # Importa o nosso gestor de chaves

# --- 2. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Compliance Ambiental", page_icon="🌿", layout="wide")

# --- 3. CARREGAR BARRA LATERAL ---
# Isto vai mostrar a chave que já inseriu, sem pedir de novo
utils.sidebar_comum()

# --- 4. VERIFICAÇÃO DE SEGURANÇA ---
# Lemos a chave diretamente da memória global
api_key = st.session_state.get("api_key", "")

if not api_key:
    st.error("🛑 **ACESSO BLOQUEADO**: A API Key não foi detetada.")
    st.info("⬅️ Por favor, insira a chave na **barra lateral esquerda** e pressione Enter.")
    st.stop() # Pára o código aqui até haver chave

# ==========================================
# DAQUI PARA BAIXO: O SEU CÓDIGO DA APP
# ==========================================
import google.generativeai as genai
# ... (Resto dos imports e lógica da app ambiente.py) ...

st.title("🌿 Módulo de Ambiente Ativo")
st.write("A chave está a funcionar e pronta a usar!")

# (Cole aqui o resto do seu código original do módulo 3...)
import streamlit as st
import google.generativeai as genai
import pypdf
from docx import Document
from io import BytesIO
import time

# --- CONFIGURAÇÃO ---
st.set_page_config(page_title="Simplex AIncA", page_icon="⚡", layout="wide")

try:
    utils.sidebar_comum()
except:
    pass

st.title("⚡ Verificação Simplex AIncA")
st.caption("Verificação de critérios de dispensa de AIA (DL 11/2023) com gestão de quota.")

# Recuperar API Key
api_key = st.session_state.get("api_key", "")
if not api_key:
    st.warning("⚠️ Configure a API Key no menu lateral.")

# --- BARRA LATERAL ESPECÍFICA ---
with st.sidebar:
    st.divider()
    st.info("Modelo: gemini-1.5-flash (Otimizado para documentos longos)")
    st.markdown("""
    **Dicas:**
    - Carregue apenas a Memória Descritiva se o PDF for > 200MB.
    - Evite peças desenhadas pesadas.
    """)

# --- FUNÇÕES ---
def get_text_from_multiple_files(file_list):
    combined_text = ""
    file_names = []
    if not file_list: return None, None

    for uploaded_file in file_list:
        try:
            reader = pypdf.PdfReader(uploaded_file)
            file_text = ""
            for page in reader.pages:
                file_text += page.extract_text() or "" 
            
            combined_text += f"\n--- FICHEIRO: {uploaded_file.name} ---\n{file_text}\n"
            file_names.append(uploaded_file.name)
        except Exception as e:
            st.error(f"Erro a ler {uploaded_file.name}: {e}")
            
    return combined_text, file_names

def create_word_docx(text, p_files, l_files):
    doc = Document()
    doc.add_heading('Relatório AIncA (Simplex)', 0)
    doc.add_paragraph(f"Projeto: {', '.join(p_files) if p_files else 'N/A'}")
    doc.add_paragraph(f"Legislação: {', '.join(l_files) if l_files else 'N/A'}")
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- ESTADO LOCAL ---
if 'analise' not in st.session_state: st.session_state['analise'] = None
if 'n_p' not in st.session_state: st.session_state['n_p'] = []
if 'n_l' not in st.session_state: st.session_state['n_l'] = []

# --- UI CENTRAL ---
col1, col2 = st.columns(2)
with col1:
    files_p = st.file_uploader("1. Projeto (Memória Descritiva)", type=["pdf"], accept_multiple_files=True)
with col2:
    files_l = st.file_uploader("2. Legislação Específica (Opcional)", type=["pdf"], accept_multiple_files=True)

if st.button("🚀 Analisar", type="primary"):
    if not api_key or not files_p:
        st.warning("Preencha a chave e carregue o Projeto.")
        st.stop()

    with st.status("A processar...", expanded=True) as status:
        status.write("A extrair texto...")
        text_p, names_p = get_text_from_multiple_files(files_p)
        text_l, names_l = get_text_from_multiple_files(files_l)
        
        # Estimativa Tokens
        total_chars = len(text_p or "") + len(text_l or "")
        estimated_tokens = total_chars / 4
        status.write(f"Tamanho detetado: ~{int(estimated_tokens)} tokens.")
        
        if estimated_tokens > 500000:
            st.warning("⚠️ Ficheiros muito grandes. A análise pode demorar.")

        # Retry Loop para Quota
        status.write("A contactar IA...")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-1.5-flash")
        
        prompt = f"""
        Atua como Consultor Ambiental.
        LEGISLAÇÃO ADICIONAL: {text_l}
        PROJETO: {text_p}
        
        TAREFA:
        Analisa a conformidade do projeto face ao DL 11/2023 (Simplex Ambiental).
        Verifica se cumpre critérios para dispensa de AIA ou se há incumprimentos óbvios.
        """

        max_retries = 3
        wait_time = 20

        for attempt in range(max_retries):
            try:
                response = model.generate_content(prompt)
                st.session_state['analise'] = response.text
                st.session_state['n_p'] = names_p
                st.session_state['n_l'] = names_l
                status.update(label="✅ Concluído!", state="complete")
                break

            except Exception as e:
                if "429" in str(e) or "Quota" in str(e):
                    if attempt < max_retries - 1:
                        status.write(f"⚠️ Limite de velocidade. A aguardar {wait_time}s (Tentativa {attempt+1}/{max_retries})...")
                        time.sleep(wait_time)
                        wait_time += 20
                    else:
                        status.update(label="❌ Falha", state="error")
                        st.error("Limite de quota excedido. Tente mais tarde.")
                else:
                    status.update(label="❌ Erro", state="error")
                    st.error(f"Erro: {e}")
                    break

# --- RESULTADOS ---
if st.session_state['analise']:
    st.divider()
    st.markdown("### 📋 Relatório")
    st.markdown(st.session_state['analise'])
    
    doc = create_word_docx(st.session_state['analise'], st.session_state['n_p'], st.session_state['n_l'])

    st.download_button("📥 Download Word", doc, "Relatorio_Simplex.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
