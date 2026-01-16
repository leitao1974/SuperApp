import sys
import os

# --- 1. CONFIGURAÇÃO DE CAMINHOS ---
current_dir = os.path.dirname(os.path.abspath(__file__))
root_dir = os.path.dirname(current_dir)
sys.path.insert(0, root_dir)

import utils
import streamlit as st
from pypdf import PdfWriter, PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
import google.generativeai as genai
import io
import time
import tempfile
from datetime import datetime

# ==========================================
# --- 2. BASE DE DADOS: CRITÉRIOS DE RIGOR (BENCHMARKS) ---
# ==========================================

# Legislação Base (Sempre verificada)
COMMON_LAWS = {
    "RJAIA (DL 151-B/2013 consolidado)": "Regime Jurídico da AIA",
    "SIMPLEX (DL 11/2023)": "Simplificação Licenciamento",
    "LUA (DL 75/2015)": "Licenciamento Único",
    "Rede Natura 2000": "DL 140/99"
}

# Benchmarks de Qualidade (O que a IA deve exigir)
SECTOR_BENCHMARKS = {
    "Energia (Eólica, Solar, Linhas)": """
    CRITÉRIOS DE RIGOR (PORTUGAL - APA/ICNF):
    1. Avifauna: O ciclo de monitorização foi ANUAL (4 estações)? Se for < 12 meses, é uma falha grave.
    2. Solar: Existe Estudo de Encandeamento (Glare)? As vedações permitem passagem de fauna (>20cm solo)?
    3. Ruído: A modelação considerou o pior cenário noturno e recetores sensíveis isolados?
    4. Cumulativos: Avaliou parques vizinhos num raio de 10km?
    """,
    "Indústria Extrativa (Minas/Pedreiras)": """
    CRITÉRIOS DE RIGOR (PORTUGAL - DGEG):
    1. PARP: O Plano de Recuperação Paisagística tem orçamento detalhado e cronograma financeiro?
    2. Vibrações: Existe estudo de uso de explosivos com sismógrafos nos edifícios vizinhos?
    3. Hidrogeologia: O cone de bombagem afeta furos de captação privados vizinhos?
    4. Poeiras: Há medidas concretas (aspersão, lavagem de rodados) ou apenas genéricas?
    """,
    "Agropecuária e Hidráulica": """
    CRITÉRIOS DE RIGOR (PORTUGAL):
    1. Efluentes: Capacidade de armazenamento para 4-6 meses (inverno)?
    2. Odores: Modelação de dispersão de odores para povoações < 500m.
    3. Água: Título de utilização hídrica (TUH) compatível com os caudais do projeto?
    """,
    "Urbanismo e Turismo": """
    CRITÉRIOS DE RIGOR:
    1. Saneamento: Ligação à rede pública garantida ou ETAR própria dimensionada?
    2. Cargas: Estudo de Tráfego considera a sazonalidade (picos de verão)?
    3. PDM: Verifica índices de impermeabilização e cérceas máximas.
    """
}

# ==========================================
# --- 3. CONFIGURAÇÃO DA PÁGINA ---
# ==========================================
st.set_page_config(
    page_title="Auditor EIA Pro (Rigor)", 
    page_icon="⚖️", 
    layout="wide"
)

# Estilo para modo "Auditor Rigoroso"
st.markdown("""
<style>
    .stButton>button { background-color: #8B0000; color: white; border-radius: 5px; font-weight: bold; }
    .stSuccess { border-left: 5px solid #228B22; }
    .stError { border-left: 5px solid #8B0000; }
</style>
""", unsafe_allow_html=True)

# --- 4. BARRA LATERAL (Setup) ---
try:
    utils.sidebar_comum()
except:
    pass

st.title("⚖️ Auditor EIA Pro (Análise Crítica & Benchmarking)")
st.markdown("""
**Módulo de Análise de Conformidade e Lacunas.**
Este sistema cruza o Processo EIA com a legislação nacional e benchmarks de boas práticas para detetar **falhas, omissões e erros fatais**.
""")

# Recuperar API Key
api_key = st.session_state.get("api_key", "")
if not api_key:
    st.warning("⚠️ **Atenção:** API Key não detetada. Por favor insira-a no menu lateral esquerdo.")
    st.stop()

# ==========================================
# --- 5. CONFIGURAÇÃO DA AUDITORIA ---
# ==========================================

with st.sidebar:
    st.divider()
    st.header("⚙️ Configuração da Auditoria")
    
    # 1. Seleção de Modelo
    def get_available_models(key):
        try:
            genai.configure(api_key=key)
            return [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        except:
            return ["models/gemini-1.5-pro", "models/gemini-1.5-flash"] # Fallback

    opcoes_modelos = get_available_models(api_key)
    # Preferência pelo PRO para raciocínio complexo, ou FLASH para volume
    idx_padrao = 0
    for i, m in enumerate(opcoes_modelos):
        if "pro" in m.lower() and "1.5" in m.lower(): idx_padrao = i; break

    selected_model = st.selectbox("Motor de Análise:", opcoes_modelos, index=idx_padrao)

    # 2. Tipologia do Projeto (Define o Benchmark)
    st.markdown("### 🏗️ Tipologia do Projeto")
    project_type = st.selectbox(
        "Selecione o setor para carregar os critérios de exigência:",
        ["Outra Tipologia"] + list(SECTOR_BENCHMARKS.keys())
    )
    
    # Carregar o texto do benchmark correspondente
    active_benchmark = SECTOR_BENCHMARKS.get(project_type, "Critérios Gerais de Boa Prática em EIA.")
    
    with st.expander("Ver Critérios Ativos"):
        st.caption(active_benchmark)

# ==========================================
# --- 6. FUNÇÕES CORE ---
# ==========================================

def merge_pdfs_to_temp(uploaded_files):
    merger = PdfWriter()
    for uploaded_file in uploaded_files:
        merger.append(uploaded_file)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        merger.write(tmp)
        tmp_path = tmp.name
    return tmp_path

def analyze_large_document(merged_pdf_path, prompt_instructions, benchmark_text, laws_dict, key, model_name):
    genai.configure(api_key=key)
    status_msg = st.empty()
    status_msg.info("📤 A enviar processo para a Google Cloud (File API)...")
    
    processo_file = None
    try:
        # 1. Upload
        processo_file = genai.upload_file(path=merged_pdf_path, display_name="Processo EIA")
        
        # 2. Polling
        status_msg.info("⚙️ A indexar volume de dados (aguarde 10-20s)...")
        while processo_file.state.name == "PROCESSING":
            time.sleep(2)
            processo_file = genai.get_file(processo_file.name)
        
        if processo_file.state.name == "FAILED":
            raise ValueError("Falha no processamento do ficheiro pela Google.")

        status_msg.success(f"✅ Indexação concluída. A iniciar Auditoria Crítica ({model_name})...")

        # 3. Montagem do Prompt Complexo
        model = genai.GenerativeModel(model_name)
        
        laws_str = "\n".join([f"- {k}: {v}" for k, v in laws_dict.items()])
        
        full_prompt = [
            prompt_instructions,
            "\n=== QUADRO LEGISLATIVO A CUMPRIR ===\n",
            laws_str,
            "\n=== BENCHMARKS DE EXIGÊNCIA TÉCNICA (NÃO IGNORAR) ===\n",
            "O projeto DEVE ser comparado com estes standards nacionais:",
            benchmark_text,
            "\n=== INSTRUÇÃO FINAL ===\n",
            "Analisa o documento em anexo. Sê implacável na procura de erros. Cita sempre a página.",
            processo_file
        ]

        # 4. Geração (Timeout alto para docs grandes)
        response = model.generate_content(full_prompt, request_options={"timeout": 600})
        
        status_msg.empty()
        return response.text

    finally:
        if processo_file:
            try: genai.delete_file(processo_file.name)
            except: pass

def create_docx(text, p_type):
    doc = Document()
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    
    title = doc.add_heading('RELATÓRIO DE AUDITORIA EIA', 0)
    title.alignment = 1
    doc.add_paragraph(f"Tipologia: {p_type} | Data: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("---")
    
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        
        if line.startswith('## '): 
            h = doc.add_heading(line.replace('##', '').strip(), 1)
            h.style.font.color.rgb = RGBColor(139, 0, 0) # Dark Red
        elif line.startswith('### '): 
            doc.add_heading(line.replace('###', '').strip(), 2)
        elif line.startswith('- ') or line.startswith('* '): 
            doc.add_paragraph(line[2:], style='List Bullet')
        else: 
            doc.add_paragraph(line)
            
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# ==========================================
# --- 7. INTERFACE PRINCIPAL ---
# ==========================================

uploaded_files = st.file_uploader(
    "Carregar Processo EIA (Tomo I, RNT, Anexos - Até 2GB)", 
    type=['pdf'], 
    accept_multiple_files=True
)

# --- INSTRUÇÕES DO AUDITOR (PERSONA) ---
instructions_audit = f"""
Atua como um **Auditor Sénior da Agência Portuguesa do Ambiente (APA)**.
A tua missão NÃO é resumir o documento, mas sim encontrar **FALHAS, OMISSÕES e INCONSISTÊNCIAS**.

Tipologia do Projeto: {project_type}

ESTRUTURA DA RESPOSTA (Markdown):

## 1. CONFORMIDADE ADMINISTRATIVA E LEGAL
   - O RNT cumpre o RJAIA? É claro para a população?
   - O projeto respeita as condicionantes (REN, RAN, Domínio Hídrico)? Cita evidências.
   - O DL 11/2023 (Simplex) foi bem aplicado?

## 2. ANÁLISE CRÍTICA VS BENCHMARKS
   - Compara o EIA com os "Benchmarks de Exigência" fornecidos. O projeto cumpre os standards nacionais?
   - **Estudo de Alternativas:** Foi real ou apenas para justificar a escolha prévia?
   - **Dados de Base:** Os dados (tráfego, ruído, fauna) são atuais (< 2 anos) ou desatualizados?

## 3. IDENTIFICAÇÃO DE "FATAL FLAWS" (ERROS GRAVES)
   - Lista pontos que inviabilizam o projeto ou requerem alterações profundas.
   - Ex: Construção em zona proibida, falta de água assegurada, perigo para saúde pública.

## 4. IMPACTES SUBVALORIZADOS PELO PROMOTOR
   - Onde é que o EIA diz "Impacte Pouco Significativo" mas tu, como perito, discordas?
   - As Medidas de Minimização são vagas (ex: "boas práticas") ou concretas?

## 5. PARECER TÉCNICO E PEDIDO DE ELEMENTOS
   - O estudo permite decidir? Ou é necessário pedir "Elementos Adicionais" (Aditamento)?
   - O que falta entregar?

REGRAS:
- Fundamenta sempre com **REFERÊNCIA À PÁGINA** do PDF (ex: "Ref: Pág. 45, Tomo I").
- Sê rigoroso, técnico e direto.
"""

if st.button("🚀 EXECUTAR AUDITORIA TÉCNICA", type="primary", use_container_width=True):
    if not uploaded_files:
        st.error("⚠️ Carregue os ficheiros do processo.")
    else:
        with st.status("🕵️‍♂️ A realizar Auditoria de Conformidade...", expanded=True) as status:
            
            status.write("📚 A consolidar volumes do processo...")
            temp_path = merge_pdfs_to_temp(uploaded_files)
            
            try:
                # Chama a função de análise com os novos parâmetros de inteligência
                res = analyze_large_document(
                    temp_path, 
                    instructions_audit, 
                    active_benchmark,
                    COMMON_LAWS,
                    api_key, 
                    selected_model
                )
                
                status.update(label="✅ Auditoria Concluída!", state="complete")
                
                st.divider()
                
                # Exibição do Relatório
                if "🚨" in res:
                    st.error(res)
                else:
                    st.subheader("📋 Parecer Técnico da IA")
                    st.markdown(res)
                    
                    # Download Word
                    doc_file = create_docx(res, project_type)
                    st.download_button(
                        label="📥 Baixar Parecer Técnico (DOCX)", 
                        data=doc_file, 
                        file_name=f"Auditoria_EIA_{project_type.split()[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
            except Exception as e:
                status.update(label="❌ Erro na Auditoria", state="error")
                st.error(f"Detalhe do erro: {e}")
                
            finally:
                try: os.remove(temp_path)
                except: pass

